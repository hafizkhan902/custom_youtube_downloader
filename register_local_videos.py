import os
import json
import sys
import subprocess
import platform

# 1. Ensure dependencies are installed
try:
    import yt_dlp
    import docx
    import cv2
except ImportError as e:
    print(f"Missing dependency: {e.name}. Installing...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "yt-dlp", "python-docx", "opencv-python"])
    import yt_dlp
    import docx
    import cv2

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Cross-platform output path resolution
if platform.system() == 'Windows':
    output_dir = os.path.join(os.path.expanduser('~'), 'Downloads', 'hafiz')
elif platform.system() == 'Darwin':
    output_dir = os.path.join(os.path.expanduser('~'), 'Desktop', 'Vibe_Download')
else:
    output_dir = "/Users/hafiz/.gemini/antigravity/scratch/trimmed_clips"

registry_path = os.path.join(output_dir, "registry.json")
doc_path = os.path.join(output_dir, "youtube_video_links.docx")

if not os.path.exists(output_dir):
    print(f"Error: Output directory does not exist at {output_dir}")
    sys.exit(1)

# 2. Load existing registry
registry = []
if os.path.exists(registry_path):
    try:
        with open(registry_path, "r") as rf:
            registry = json.load(rf)
        print(f"Loaded existing registry with {len(registry)} entries.")
    except Exception as e:
        print(f"Warning: Could not read existing registry: {e}")

# Keep track of already registered filenames (from registry)
registered_filenames = {item['filename'] for item in registry}

# 3. Scan output directory for any local video files starting with 'clip_'
all_local_files = [f for f in os.listdir(output_dir) if f.startswith("clip_") and f.lower().endswith(('.mp4', '.mkv', '.webm'))]

# Filter out files that are already registered
unregistered_files = [f for f in all_local_files if f not in registered_filenames]

if not unregistered_files:
    print("No unregistered local video files found.")
else:
    print(f"Found {len(unregistered_files)} unregistered local video clips.")

    # Deduplicate files that have the same content (e.g. clip_11_Letting_go...mkv and .mp4.mkv)
    # Group them by base title and size
    unregistered_groups = {}
    for f_name in unregistered_files:
        # Extract title words
        parts = f_name.split('_')
        title_parts = []
        for p in parts[2:] if len(parts) > 2 else []:
            if p:
                # Remove common extensions from parts
                p_clean = p.replace(".mp4.mkv", "").replace(".mp4.webm", "").replace(".mp4", "").replace(".mkv", "").replace(".webm", "")
                if p_clean:
                    title_parts.append(p_clean)
        
        normalized_title = " ".join(title_parts).strip().lower()
        size = os.path.getsize(os.path.join(output_dir, f_name))
        
        key = (normalized_title, size)
        if key not in unregistered_groups:
            unregistered_groups[key] = []
        unregistered_groups[key].append(f_name)
    
    unique_unregistered = []
    for key, files in unregistered_groups.items():
        files_sorted = sorted(files, key=len, reverse=True)
        chosen = files_sorted[0]
        unique_unregistered.append(chosen)
        
        for extra in files_sorted[1:]:
            extra_path = os.path.join(output_dir, extra)
            try:
                os.remove(extra_path)
                print(f"Deleted duplicate file: {extra}")
            except Exception as e:
                print(f"Failed to delete duplicate {extra}: {e}")

    print(f"Unique unregistered files to process: {len(unique_unregistered)}")

    # Process each unique unregistered file
    for idx, f_name in enumerate(unique_unregistered):
        next_idx = len(registry) + 1
        print(f"\n[{idx+1}/{len(unique_unregistered)}] Processing: {f_name}")
        
        # Extract title from filename
        parts = f_name.split('_')
        title_parts = []
        for p in parts[2:] if len(parts) > 2 else []:
            if p:
                p_clean = p.replace(".mp4.mkv", "").replace(".mp4.webm", "").replace(".mp4", "").replace(".mkv", "").replace(".webm", "")
                if p_clean:
                    title_parts.append(p_clean)
        
        title_clean = " ".join(title_parts)
        title_query = " ".join(title_clean.replace("_", " ").split())
        
        # Format the clean filename prefix
        clean_filename_title = "_".join(title_parts).strip("_")
        new_filename = f"clip_{next_idx:02d}_{clean_filename_title}.mp4"
        
        input_path = os.path.join(output_dir, f_name)
        temp_output_path = os.path.join(output_dir, f"temp_convert_{next_idx}.mp4")
        final_output_path = os.path.join(output_dir, new_filename)
        
        print("  Converting to standard MP4 using ffmpeg...")
        command = [
            'ffmpeg', '-y',
            '-i', input_path,
            '-c:v', 'libx264', '-preset', 'fast',
            '-c:a', 'aac',
            temp_output_path
        ]
        
        success = False
        try:
            result = subprocess.run(command, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
            if result.returncode == 0:
                success = True
                print("  Conversion successful.")
            else:
                print(f"  FFmpeg conversion failed:\n{result.stderr}")
        except Exception as e:
            print(f"  Error running FFmpeg: {e}")
            
        if success and os.path.exists(temp_output_path) and os.path.getsize(temp_output_path) > 0:
            try:
                os.remove(input_path)
            except Exception as e:
                print(f"  Warning: Could not delete original file {f_name}: {e}")
                
            if os.path.exists(final_output_path):
                os.remove(final_output_path)
            os.rename(temp_output_path, final_output_path)
            print(f"  Saved as: {new_filename}")
        else:
            print("  Skipping registration due to conversion failure.")
            if os.path.exists(temp_output_path):
                try:
                    os.remove(temp_output_path)
                except:
                    pass
            continue
            
        # Get video duration from the local file
        duration = 30.0
        try:
            cap = cv2.VideoCapture(final_output_path)
            if cap.isOpened():
                fps = cap.get(cv2.CAP_PROP_FPS)
                frame_count = cap.get(cv2.CAP_PROP_FRAME_COUNT)
                if fps > 0:
                    duration = frame_count / fps
                cap.release()
        except Exception as e:
            print(f"  Warning: Could not determine file duration: {e}")
            
        # Fetch metadata from YouTube
        search_query = f"ytsearch1:TED talk {title_query}"
        print(f"  Searching YouTube for: 'TED talk {title_query}'")
        
        vid_title = title_query
        vid_url = ""
        
        try:
            ydl_opts = {'format': '160', 'quiet': True}
            with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                res = ydl.extract_info(search_query, download=False)
                if 'entries' in res and res['entries']:
                    vid_info = res['entries'][0]
                    vid_url = f"https://www.youtube.com/watch?v={vid_info['id']}"
                    vid_title = vid_info.get('title', title_query)
                    print(f"  Matched YouTube video: '{vid_title}' ({vid_url})")
                else:
                    print(f"  No YouTube match found. Using filename title.")
        except Exception as e:
            print(f"  Failed to retrieve YouTube metadata: {e}")
            
        # Add to registry
        registry.append({
            "clip_index": next_idx,
            "filename": new_filename,
            "title": vid_title,
            "youtube_url": vid_url,
            "start_seconds": 0.0,
            "end_seconds": round(duration, 2),
            "timestamp_start": "00:00",
            "timestamp_end": f"{int(duration // 60):02d}:{int(duration % 60):02d}"
        })
        
        with open(registry_path, "w") as rf:
            json.dump(registry, rf, indent=2)

# Ensure sorted registry by index
registry = sorted(registry, key=lambda x: x.get('clip_index', 0))
with open(registry_path, "w") as rf:
    json.dump(registry, rf, indent=2)
print(f"\nRegistry successfully saved to: {registry_path}")

# 4. Generate the Google Docs-compatible Word Document
print("Generating Word Document...")
try:
    doc = Document()
    
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("YouTube Video Links & Timestamps")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(18)
    title_run.bold = True
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    intro_p = doc.add_paragraph()
    intro_p.add_run(
        "This document lists all the steady speaker close-up clips downloaded onto your local machine. "
        "To access the online YouTube sources directly, click any link below. "
        "You can upload this file to Google Drive/Google Docs for cloud storage and sharing."
    ).italic = True
    
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '#'
    hdr_cells[1].text = 'Video Title'
    hdr_cells[2].text = 'YouTube Link'
    hdr_cells[3].text = 'Timestamps'
    hdr_cells[4].text = 'Filename'
    
    for cell in hdr_cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.name = 'Arial'
                run.font.size = Pt(11)
                
    for item in registry:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item['clip_index'])
        row_cells[1].text = item['title']
        row_cells[2].text = item['youtube_url']
        row_cells[3].text = f"{item['timestamp_start']} - {item['timestamp_end']}"
        row_cells[4].text = item['filename']
        
        for cell in row_cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)
                    
    doc.save(doc_path)
    print(f"Document successfully created at: {doc_path}")
except Exception as e:
    print(f"Error generating Word Document: {e}")
