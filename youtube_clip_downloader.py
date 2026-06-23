import yt_dlp
import cv2
import numpy as np
import os
import sys
import time
import subprocess
import argparse
import webbrowser
import select
import urllib.parse
import json
import platform

# Timed input helper function for slow buffer warnings
def input_with_timeout(prompt, timeout=10):
    print(prompt, end='', flush=True)
    try:
        ready, _, _ = select.select([sys.stdin], [], [], timeout)
        if ready:
            return sys.stdin.readline().strip().lower()
        else:
            print()
            return ""
    except Exception:
        # Fallback to standard blocking input if select is not supported
        try:
            return input().strip().lower()
        except:
            return ""

# Helper to check custom face count criteria
def check_face_count(count, criteria):
    if isinstance(criteria, int):
        return count == criteria
    if criteria == '>=1':
        return count >= 1
    if criteria == '0' or criteria == 0:
        return count == 0
    try:
        if criteria.startswith('>='):
            return count >= int(criteria[2:])
        elif criteria.startswith('<='):
            return count <= int(criteria[2:])
        else:
            return count == int(criteria)
    except:
        return count >= 1

# Q&A Wizard configuration flow
def run_qa_wizard():
    print("="*60)
    print("      YOUTUBE VIDEO CLIP DOWNLOADER & ANALYZER WIZARD")
    print("="*60)
    
    # 1. Download videos option
    dl_input = input("Download trimmed video clips to local disk? (y/n, default: y): ").strip().lower()
    download_videos = False if dl_input == 'n' else True
    
    # 2. Topic/Query
    query = input("Enter YouTube search query (default: 'TED talk'): ").strip()
    if not query:
        query = "TED talk"
        
    # 3. CC Filter
    cc_input = input("Filter by Creative Commons license only? (y/n, default: y): ").strip().lower()
    cc = False if cc_input == 'n' else True
    
    # 4. Clip Count
    count_input = input("How many clips do you want to download? (default: 50): ").strip()
    try:
        count = int(count_input) if count_input else 50
    except ValueError:
        count = 50
        
    # 5. Duration
    min_input = input("Minimum duration of the clip in seconds? (default: 22): ").strip()
    try:
        min_len = float(min_input) if min_input else 22.0
    except ValueError:
        min_len = 22.0
        
    max_input = input("Maximum duration of the clip in seconds? (default: 30): ").strip()
    try:
        max_len = float(max_input) if max_input else 30.0
    except ValueError:
        max_len = 30.0
        
    # 6. Analysis Mode
    print("\nChoose visual analysis mode:")
    print("  [1] Steady Single Speaker close-up (exactly 1 face detected, default)")
    print("  [2] Two Speakers / Dialogue (exactly 2 faces detected)")
    print("  [3] Any Face present (at least 1 face detected)")
    print("  [4] Generic / No face detection (takes first uncut shot of duration)")
    print("  [5] Others (Configure custom face count, frequency, and size)")
    mode_input = input("Choose option [1-5] (default: 1): ").strip()
    
    mode = "single-speaker"
    custom_settings = {}
    if mode_input == '2':
        mode = "two-speakers"
    elif mode_input == '3':
        mode = "any-face"
    elif mode_input == '4':
        mode = "generic"
    elif mode_input == '5':
        mode = "others"
        print("\n--- Custom Mode Settings ---")
        target_faces = input("Enter target face count (e.g., '0', '3', or '>=1') [default: >=1]: ").strip()
        if not target_faces:
            target_faces = ">=1"
            
        freq_input = input("Enter minimum face presence frequency (0.0 to 1.0) [default: 0.8]: ").strip()
        try:
            freq = float(freq_input) if freq_input else 0.8
        except ValueError:
            freq = 0.8
            
        size_input = input("Enter minimum face size relative to frame width (0.0 to 1.0) [default: 0.08]: ").strip()
        try:
            size = float(size_input) if size_input else 0.08
        except ValueError:
            size = 0.08
            
        custom_settings = {
            "target_faces": target_faces,
            "min_frequency": freq,
            "min_size": size
        }
    else:
        mode = "single-speaker"
        
    # 7. Output folder (resolved dynamically by OS)
    default_folder = "trimmed_clips"
    if platform.system() == 'Windows':
        default_folder = os.path.join(os.path.expanduser('~'), 'Downloads', 'hafiz')
    elif platform.system() == 'Darwin':
        default_folder = os.path.join(os.path.expanduser('~'), 'Desktop', 'Vibe_Download')
        
    folder = input(f"\nEnter output folder name (default: '{default_folder}'): ").strip()
    if not folder:
        folder = default_folder
        
    return {
        "download_videos": download_videos,
        "query": query,
        "cc": cc,
        "count": count,
        "min_len": min_len,
        "max_len": max_len,
        "mode": mode,
        "custom_settings": custom_settings,
        "output_dir": folder
    }

# Parse command-line options
parser = argparse.ArgumentParser(description="Download dynamic clips from YouTube.")
parser.add_argument("--login", action="store_true", help="Log in to YouTube using Google Device Activation.")
parser.add_argument("--cookies", type=str, choices=["chrome", "safari", "firefox", "edge", "opera"],
                    help="Extract cookies from the specified desktop browser to bypass bot detection.")
# Dynamic command-line parameters
parser.add_argument("--download", dest="download", action="store_true", help="Download video files.")
parser.add_argument("--no-download", dest="download", action="store_false", help="Do not download videos (Metadata Mode).")
parser.set_defaults(download=True)
parser.add_argument("--query", "-q", type=str, help="YouTube search query.")
parser.add_argument("--cc", type=str, choices=["y", "n"], help="Filter by Creative Commons (y/n).")
parser.add_argument("--count", "-n", type=int, help="Target number of clips to download.")
parser.add_argument("--min-len", "--min-duration", type=float, help="Minimum duration of trimmed clips.")
parser.add_argument("--max-len", "--max-duration", type=float, help="Maximum duration of trimmed clips.")
parser.add_argument("--mode", "-m", type=str, choices=["single-speaker", "two-speakers", "any-face", "generic", "others"],
                    help="Visual analysis mode.")
parser.add_argument("--output-dir", "-o", type=str, help="Output directory path.")
parser.add_argument("--custom-target-faces", type=str, default=">=1", help="Custom target face count (e.g. '0', '3', '>=1').")
parser.add_argument("--custom-min-frequency", type=float, default=0.8, help="Custom min face frequency.")
parser.add_argument("--custom-min-size", type=float, default=0.08, help="Custom min face size.")
args = parser.parse_args()

# Handle Google authentication opening
if args.login:
    print("\n[Auth] Opening Google Device Activation in your default browser...")
    webbrowser.open("https://google.com/device")
    print("[Auth] Please enter the code displayed below on that webpage to log in.\n")

# Ensure python-docx is installed for document generation
try:
    import docx
except ImportError:
    print("Installing python-docx library for document creation...")
    try:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        import docx
    except Exception as e:
        print(f"Warning: Could not install python-docx: {e}. The document will not be generated automatically.")

# Check if arguments were supplied to run in CLI non-interactive mode
is_cli_run = (args.query is not None or 
              args.cc is not None or 
              args.count is not None or 
              args.min_len is not None or 
              args.max_len is not None or 
              args.mode is not None or 
              args.output_dir is not None)

if is_cli_run:
    download_videos = args.download
    query = args.query if args.query else "TED talk"
    cc = False if args.cc == 'n' else True
    target_count = args.count if args.count else 50
    min_len = args.min_len if args.min_len else 22.0
    max_len = args.max_len if args.max_len else 30.0
    mode = args.mode if args.mode else "single-speaker"
    folder_name = args.output_dir if args.output_dir else "trimmed_clips"
    custom_settings = {
        "target_faces": args.custom_target_faces,
        "min_frequency": args.custom_min_frequency,
        "min_size": args.custom_min_size
    }
else:
    config = run_qa_wizard()
    download_videos = config["download_videos"]
    query = config["query"]
    cc = config["cc"]
    target_count = config["count"]
    min_len = config["min_len"]
    max_len = config["max_len"]
    mode = config["mode"]
    custom_settings = config["custom_settings"]
    folder_name = config["output_dir"]

# Resolve output path based on cross-platform rules
if os.path.isabs(folder_name):
    output_dir = folder_name
else:
    if platform.system() == 'Windows':
        output_dir = os.path.join(os.path.expanduser('~'), 'Downloads', 'hafiz', folder_name)
    elif platform.system() == 'Darwin':
        output_dir = os.path.join(os.path.expanduser('~'), 'Desktop', 'Vibe_Download', folder_name)
    else:
        output_dir = os.path.join(os.getcwd(), folder_name)

# Print Configuration Summary
print("\n" + "="*60)
print("                  STARTING SCANNING & DOWNLOAD")
print("="*60)
print(f"Download Video Clips:  {'Yes' if download_videos else 'No (Metadata Mode)'}")
print(f"Search Query:          {query}")
print(f"Creative Commons:      {'Yes' if cc else 'No'}")
print(f"Target Clip Count:     {target_count}")
print(f"Clip Duration Range:   {min_len}s - {max_len}s")
print(f"Visual Analysis Mode:  {mode}")
if mode == 'others':
    print(f"  - Target Face Count: {custom_settings.get('target_faces')}")
    print(f"  - Min Frequency:     {custom_settings.get('min_frequency')}")
    print(f"  - Min Face Size:     {custom_settings.get('min_size')}")
print(f"Output Directory:      {output_dir}")
print("="*60 + "\n")

# Setup directories
os.makedirs(output_dir, exist_ok=True)

# Load face cascade
face_cascade_path = cv2.data.haarcascades + 'haarcascade_frontalface_default.xml'
face_cascade = cv2.CascadeClassifier(face_cascade_path)
if face_cascade.empty():
    print("Error: Could not load face cascade.")
    sys.exit(1)

# Build search URL
if cc:
    search_query = f"https://www.youtube.com/results?search_query={urllib.parse.quote_plus(query)}&sp=EgIwAQ%3D%3D"
else:
    search_query = f"https://www.youtube.com/results?search_query={urllib.parse.quote_plus(query)}"

ydl_opts = {
    'quiet': True,
    'extract_flat': True,
    'playlist_items': '1-120',  # Extract up to 120 search results
}

# Apply auth settings to search request
if args.login:
    ydl_opts['username'] = 'oauth2'
    ydl_opts['password'] = ''
if args.cookies:
    ydl_opts['cookiesfrombrowser'] = (args.cookies,)

print(f"Searching YouTube: '{search_query}'...")
video_urls = []
with yt_dlp.YoutubeDL(ydl_opts) as ydl:
    try:
        result = ydl.extract_info(search_query, download=False)
        if 'entries' in result:
            for entry in result['entries']:
                video_urls.append(f"https://www.youtube.com/watch?v={entry['id']}")
    except Exception as e:
        print(f"Error fetching search results: {e}")
        sys.exit(1)

print(f"Found {len(video_urls)} candidate videos. Starting analysis...")

registry = []
downloaded_urls = set()

# Load existing registry
registry_path = os.path.join(output_dir, "registry.json")
if os.path.exists(registry_path):
    try:
        with open(registry_path, "r") as rf:
            registry = json.load(rf)
        print(f"Loaded existing registry with {len(registry)} clips.")
    except Exception as e:
        print(f"Warning: Could not load existing registry: {e}")

# Reconcile local files in the directory with registry.json
local_files = [f for f in os.listdir(output_dir) if f.startswith("clip_") and f.endswith(".mp4")]
registered_filenames = {item['filename'] for item in registry}
unregistered_files = [f for f in local_files if f not in registered_filenames]

if unregistered_files:
    print(f"\nFound {len(unregistered_files)} unregistered local video clips. Reconstructing metadata...")
    for f_name in unregistered_files:
        # Reconstruct search query from filename (e.g. clip_01_Do_schools_kill_creativity.mp4)
        parts = f_name.split('_')
        title_words = [p for p in parts[2:] if p]
        title_clean = " ".join(title_words).replace(".mp4", "").strip()
        title_query = " ".join(title_clean.replace("_", " ").split())
        
        print(f"  Registering local file: {f_name}...")
        search_q = f"ytsearch1:{query} {title_query}"
        try:
            with yt_dlp.YoutubeDL({'format': '160', 'quiet': True}) as ydl:
                res = ydl.extract_info(search_q, download=False)
                if 'entries' in res and res['entries']:
                    vid_info = res['entries'][0]
                    vid_url = f"https://www.youtube.com/watch?v={vid_info['id']}"
                    vid_title = vid_info.get('title', title_query)
                    
                    registry.append({
                        "clip_index": len(registry) + 1,
                        "filename": f_name,
                        "title": vid_title,
                        "youtube_url": vid_url,
                        "start_seconds": 0.0,
                        "end_seconds": max_len,
                        "timestamp_start": "00:00",
                        "timestamp_end": f"{int(max_len // 60):02d}:{int(max_len % 60):02d}"
                    })
                    print(f"  Successfully registered: '{vid_title}' -> {vid_url}")
        except Exception as e:
            print(f"  Failed to retrieve metadata for {f_name}: {e}")
            
    # Save the reconciled registry.json
    with open(registry_path, "w") as rf:
        json.dump(registry, rf, indent=2)
    print("Reconciliation complete. Registry file updated.")

downloaded_urls = {item['youtube_url'] for item in registry}
downloaded_count = len(registry)

# Initialize persistent yt-dlp instances to speed up execution
scan_opts = {
    'format': '160',
    'quiet': True,
    'no_warnings': True,
    'youtube_include_dash_manifest': False,
    'youtube_include_hls_manifest': False,
    'nocheckcertificate': True,
}

download_opts = {
    'format': 'bestvideo[height<=1080]+bestaudio/best[height<=1080]/best',
    'quiet': True,
    'force_keyframes_at_cuts': True,
    'http_chunk_size': 1048576,  # Prevent YouTube throttling
    'nocheckcertificate': True,
    'merge_output_format': 'mp4',  # Ensure output is merged to MP4 container
    'remux_video': 'mp4',          # Force ffmpeg to output as standard MP4 container
}

if args.login:
    scan_opts['username'] = 'oauth2'
    scan_opts['password'] = ''
    download_opts['username'] = 'oauth2'
    download_opts['password'] = ''

if args.cookies:
    scan_opts['cookiesfrombrowser'] = (args.cookies,)
    download_opts['cookiesfrombrowser'] = (args.cookies,)

ydl_scanner = yt_dlp.YoutubeDL(scan_opts)
ydl_downloader = yt_dlp.YoutubeDL(download_opts)

scan_durations = []

for idx, url in enumerate(video_urls):
    if downloaded_count >= target_count:
        print(f"Reached target of {target_count} downloaded videos.")
        break
        
    if url in downloaded_urls:
        print(f"Skipping: Video {url} has already been downloaded (found in local registry).")
        continue
        
    print(f"\n--- Analyzing video {idx + 1}/{len(video_urls)}: {url} ---")
    
    try:
        info = ydl_scanner.extract_info(url, download=False)
        stream_url = info.get('url')
        video_title = info.get('title', f"video_{idx}")
        license = info.get('license', '') or ''
        desc = info.get('description', '') or ''
        
        # Verify it is Creative Commons if CC filter is enabled
        if cc:
            is_cc = ('creative commons' in license.lower()) or ('creative commons' in desc.lower()) or ('cc by' in desc.lower())
            if not is_cc:
                print(f"Skipping: Not a Creative Commons video (License: '{license}')")
                continue
            else:
                print(f"Verified CC: {video_title}")
        else:
            print(f"Analyzing: {video_title}")
    except Exception as e:
        print(f"Error fetching stream info: {e}")
        continue
        
    # Dynamic scanning timeout based on previous scans
    if len(scan_durations) >= 3:
        avg_time = sum(scan_durations) / len(scan_durations)
        scan_timeout = max(10.0, avg_time * 1.8)
    else:
        scan_timeout = 25.0
        
    analysis_start_time = time.time()
    scan_was_slow = False
    
    # Open capture
    cap = cv2.VideoCapture(stream_url)
    if not cap.isOpened():
        print("Error: Could not open stream.")
        continue
        
    fps = cap.get(cv2.CAP_PROP_FPS)
    if not fps or fps < 10:
        fps = 30.0
        
    max_scan_seconds = 180.0
    frame_step = int(fps / 2) if fps >= 2 else 1  # 2 frames per second
    
    prev_hist = None
    shots = []
    current_shot_start = 0.0
    
    frame_idx = 0
    max_frame = int(max_scan_seconds * fps)
    
    # Scan for cuts
    print("  Scanning video for camera cuts...")
    while frame_idx < max_frame:
        cap.set(cv2.CAP_PROP_POS_FRAMES, frame_idx)
        ret, frame = cap.read()
        if not ret:
            break
            
        timestamp = frame_idx / fps
        gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
        
        hist = cv2.calcHist([gray], [0], None, [256], [0, 256])
        cv2.normalize(hist, hist)
        
        if prev_hist is not None:
            corr = cv2.compareHist(prev_hist, hist, cv2.HISTCMP_CORREL)
            if corr < 0.85:
                duration = timestamp - current_shot_start
                if duration >= min_len:
                    shots.append((current_shot_start, timestamp))
                current_shot_start = timestamp
                
        prev_hist = hist
        frame_idx += frame_step
        
        # Timed non-blocking buffer warning prompt
        if time.time() - analysis_start_time > scan_timeout:
            prompt = "\n  This video buffer streaming average is low .. do you want to wait ? or want to skip ? (Type 's' to skip, 'w' or press enter to wait): "
            user_choice = input_with_timeout(prompt, timeout=10)
            if user_choice in ['s', 'skip']:
                print("  Skipping this video as requested.")
                scan_was_slow = True
                break
            else:
                print("  Continuing to scan...")
                analysis_start_time = time.time()  # extend timeout
            
        # Display progress bar for scanning
        if (frame_idx // frame_step) % 5 == 0:
            percent = min(100, int((frame_idx / max_frame) * 100))
            bar_length = 20
            filled_length = int(bar_length * percent // 100)
            bar = '█' * filled_length + '░' * (bar_length - filled_length)
            sys.stdout.write(f"\r  🔍 [Scanning]  [{bar}]  {percent:3d}%  |  Time: {timestamp:5.1f}s / {max_scan_seconds:.0f}s  |  Uncut clips found: {len(shots):2d}")
            sys.stdout.flush()
            
    # Add trailing shot if valid
    if not scan_was_slow:
        end_timestamp = frame_idx / fps
        if end_timestamp - current_shot_start >= min_len:
            shots.append((current_shot_start, end_timestamp))
            
        bar = '█' * 20
        sys.stdout.write(f"\r  🔍 [Scanning]  [{bar}]  100%  |  Time: {max_scan_seconds:5.1f}s / {max_scan_seconds:.0f}s  |  Uncut clips found: {len(shots):2d}\n")
        sys.stdout.flush()
        
    if scan_was_slow:
        cap.release()
        continue
        
    # Initialize valid shots list
    valid_shots = []
    
    # Visual analysis
    if mode == 'generic':
        # Generic mode bypasses face check; every valid duration shot matches
        valid_shots = shots
    else:
        if shots:
            print(f"  Verifying visual criteria in the {len(shots)} uncut candidate clips...")
            
        for shot_idx, (s_start, s_end) in enumerate(shots):
            # Timed warning check inside visual analysis phase
            if time.time() - analysis_start_time > scan_timeout + 8.0:
                prompt = "\n  This video buffer streaming average is low .. do you want to wait ? or want to skip ? (Type 's' to skip, 'w' or press enter to wait): "
                user_choice = input_with_timeout(prompt, timeout=10)
                if user_choice in ['s', 'skip']:
                    print("  Skipping this video as requested.")
                    scan_was_slow = True
                    break
                else:
                    print("  Continuing to scan...")
                    analysis_start_time = time.time()  # extend timeout
                
            total_shots = len(shots)
            current_shot = shot_idx + 1
            percent_anal = int(100 * current_shot / total_shots) if total_shots > 0 else 100
            bar_length_anal = 15
            filled_length_anal = int(bar_length_anal * percent_anal // 100)
            bar_anal = '█' * filled_length_anal + '░' * (bar_length_anal - filled_length_anal)
            sys.stdout.write(f"\r  🧠 [Analyzing] [{bar_anal}] {percent_anal:3d}% | Shot {current_shot}/{total_shots} ({s_start:.1f}s - {s_end:.1f}s)...")
            sys.stdout.flush()
            
            # Sample 5 frames in this shot to check faces
            sample_times = np.linspace(s_start, s_end, 7)[1:-1]
            face_count = 0
            face_sizes = []
            
            for t in sample_times:
                f_num = int(t * fps)
                cap.set(cv2.CAP_PROP_POS_FRAMES, f_num)
                ret, frame = cap.read()
                if not ret:
                    continue
                gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
                faces = face_cascade.detectMultiScale(gray, 1.1, 4, minSize=(20, 20))
                
                # Check criteria match
                is_match = False
                if mode == 'single-speaker':
                    is_match = (len(faces) == 1)
                elif mode == 'two-speakers':
                    is_match = (len(faces) == 2)
                elif mode == 'any-face':
                    is_match = (len(faces) >= 1)
                elif mode == 'others':
                    is_match = check_face_count(len(faces), custom_settings.get("target_faces", ">=1"))
                    
                if is_match:
                    face_count += 1
                    img_w = frame.shape[1]
                    if len(faces) > 0:
                        w_max = max(face[2] for face in faces)
                        face_sizes.append(w_max / img_w)
                    else:
                        face_sizes.append(0.0)
                        
            face_ratio = face_count / 5.0
            avg_face_size = np.mean(face_sizes) if face_sizes else 0
            
            # Apply configured criteria limits
            req_ratio = 0.8
            req_size = 0.08
            
            if mode == 'others':
                req_ratio = custom_settings.get("min_frequency", 0.8)
                req_size = custom_settings.get("min_size", 0.08)
                if custom_settings.get("target_faces") in ['0', 0]:
                    req_size = 0.0
            elif mode in ['two-speakers', 'any-face']:
                req_size = 0.05
                
            if face_ratio >= req_ratio and (req_size == 0.0 or avg_face_size >= req_size):
                valid_shots.append((s_start, s_end))
                print(f"\n    -> MATCH FOUND: Clip {s_start:.1f}s-{s_end:.1f}s (Ratio: {face_ratio:.0%}, Size: {avg_face_size:.1%})")
                
        if shots and not scan_was_slow:
            bar_anal = '█' * 15
            sys.stdout.write(f"\r  🧠 [Analyzing] [{bar_anal}] 100% | Completed visual checks for all {len(shots)} candidates.\n")
            sys.stdout.flush()
            
    cap.release()
    
    if scan_was_slow:
        continue
        
    # Record scan duration
    elapsed = time.time() - analysis_start_time
    scan_durations.append(elapsed)
    if len(scan_durations) > 10:
        scan_durations.pop(0)
        
    if not valid_shots:
        print("No match found satisfying visual criteria in this video.")
        continue
        
    # Take the best shot
    best_shot = sorted(valid_shots, key=lambda x: x[1] - x[0], reverse=True)[0]
    s_start, s_end = best_shot
    
    if s_end - s_start > max_len:
        s_end = s_start + max_len
        
    clean_title = "".join([c if c.isalnum() else "_" for c in video_title]).strip("_")[:30]
    if not clean_title:
        clean_title = "clip"
        
    out_filename = f"clip_{downloaded_count + 1:02d}_{clean_title}.mp4"
    out_path = os.path.join(output_dir, out_filename)
    
    # We use a %(ext)s template to let yt-dlp merge/remux correctly, then resolve to mp4.
    out_filename_template = f"clip_{downloaded_count + 1:02d}_{clean_title}.%(ext)s"
    out_path_template = os.path.join(output_dir, out_filename_template)
    
    if download_videos:
        print(f"Downloading stream: {s_start:.2f}s to {s_end:.2f}s...")
        
        # Configure download range and output destination
        ydl_downloader.params['download_ranges'] = lambda info, ctx: [{'start_time': s_start, 'end_time': s_end}]
        ydl_downloader.params['outtmpl'] = {'default': out_path_template}
        try:
            ydl_downloader.download([url])
                
            if os.path.exists(out_path) and os.path.getsize(out_path) > 0:
                downloaded_count += 1
                print(f"Saved: {out_filename} (Success! Total: {downloaded_count}/{target_count})")
                registry.append({
                    "clip_index": downloaded_count,
                    "filename": out_filename,
                    "title": video_title,
                    "youtube_url": url,
                    "start_seconds": s_start,
                    "end_seconds": s_end,
                    "timestamp_start": f"{int(s_start // 60):02d}:{int(s_start % 60):02d}",
                    "timestamp_end": f"{int(s_end // 60):02d}:{int(s_end % 60):02d}"
                })
                
                with open(registry_path, "w") as rf:
                    json.dump(registry, rf, indent=2)
            else:
                print("Failed to download or write trimmed clip.")
        except Exception as e:
            print(f"Error during download and trim: {e}")
    else:
        # Metadata only (No download)
        downloaded_count += 1
        print(f"Registered (Metadata Only): '{video_title}' at {s_start:.1f}s-{s_end:.1f}s")
        registry.append({
            "clip_index": downloaded_count,
            "filename": f"[Metadata Only] {out_filename}",
            "title": video_title,
            "youtube_url": url,
            "start_seconds": s_start,
            "end_seconds": s_end,
            "timestamp_start": f"{int(s_start // 60):02d}:{int(s_start % 60):02d}",
            "timestamp_end": f"{int(s_end // 60):02d}:{int(s_end % 60):02d}"
        })
        
        with open(registry_path, "w") as rf:
            json.dump(registry, rf, indent=2)

print(f"\nCompleted! Processed {downloaded_count} clips into: {output_dir}")
print("Metadata registry saved at: " + registry_path)

# Generate Word Document if docx is installed
if 'docx' in sys.modules:
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc_path = os.path.join(output_dir, "youtube_video_links.docx")
        print("\nGenerating Word Document...")
        doc = Document()
        
        title_p = doc.add_paragraph()
        title_text = f"YouTube Video Links & Timestamps for '{query}'"
        if cc:
            title_text += " (Creative Commons)"
        title_run = title_p.add_run(title_text)
        title_run.font.name = 'Arial'
        title_run.font.size = Pt(18)
        title_run.bold = True
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        intro_p = doc.add_paragraph()
        intro_p.add_run(
            f"This document lists all the trimmed clips downloaded onto your local machine for the search topic '{query}'. "
            "To access the online YouTube sources directly, click any link below. "
            "You can upload this file to Google Drive/Google Docs for cloud storage and sharing."
        ).italic = True
        
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '#'
        hdr_cells[1].text = 'Video Title'
        hdr_cells[2].text = 'YouTube Link'
        hdr_cells[3].text = 'Timestamps'
        hdr_cells[4].text = 'Filename'
        
        for cell in hdr_cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.bold = True
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    
        for item in registry:
            row_cells = table.add_row().cells
            row_cells[0].text = str(item['clip_index'])
            row_cells[1].text = item['title']
            row_cells[2].text = item['youtube_url']
            row_cells[3].text = f"{item['timestamp_start']} - {item['timestamp_end']}"
            row_cells[4].text = item['filename']
            
            for cell in row_cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(10)
                        
        doc.save(doc_path)
        print(f"Document successfully generated at: {doc_path}")
    except Exception as e:
        print(f"Error generating Word Document: {e}")
